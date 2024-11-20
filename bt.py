import os
import tkinter
from tkinter import Tk, filedialog
import tkinter as tk
import tkinter.messagebox
from docx import Document
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import customtkinter
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import docx
import sys
from  docx.api import Document
from docx.shared import Pt, RGBColor 
from docx.oxml.ns import qn
import json
from tkinter import filedialog
import time
import webbrowser

from reportlab.lib.pagesizes import letter 
from reportlab.pdfgen import canvas

customtkinter.set_appearance_mode("System")  # Modes: "System" (standard), "Dark", "Light"
customtkinter.set_default_color_theme("blue")  # Themes: "blue" (standard), "green", "dark-blue"


class App(customtkinter.CTk):
    def __init__(self):
        super().__init__()
        self.file = ""

        # configure window
        self.title("Asquini-Encorad")
        self.geometry(f"{1080}x{720}")
        # self.iconbitmap("logo.png")

        # configure grid layout (4x4)
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure((2, 3), weight=0)
        self.grid_rowconfigure((1, 1, 1), weight=1)

        # create sidebar frame with widgets
        self.sidebar_frame = customtkinter.CTkFrame(self, width=140, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, rowspan=4, sticky="nsew")
        self.sidebar_frame.grid_rowconfigure(4, weight=1)
        
        self.logo_label = customtkinter.CTkLabel(
            self.sidebar_frame, text="Qantix", 
            font=customtkinter.CTkFont(size=24, weight="bold"),
            text_color="#DC3F3F"
            )
        self.logo_label.grid(row=0, column=0, padx=20, pady=(20, 10))
        
        self.sidebar_button_1 = customtkinter.CTkButton(self.sidebar_frame, command=self.import_file, text="Import file", fg_color="#293DA4")
        self.sidebar_button_1.grid(row=1, column=0, padx=20, pady=10)
        self.appearance_mode_label = customtkinter.CTkLabel(self.sidebar_frame, text="Appearance Mode:", anchor="w",)
        self.appearance_mode_label.grid(row=5, column=0, padx=20, pady=(10, 0))
        self.appearance_mode_optionemenu = customtkinter.CTkOptionMenu(self.sidebar_frame, values=["Light", "Dark", "System"],
                                                                       command=self.change_appearance_mode_event, fg_color="#293DA4")
        self.appearance_mode_optionemenu.grid(row=6, column=0, padx=20, pady=(10, 10))
        
        
        #create header
        self.head = customtkinter.CTkFrame(master=self)
        self.head.grid(row=0, column=1, sticky="new")

        # create main entry and button
        self.bt_option = customtkinter.CTkOptionMenu(
            master=self.head, 
            values=["Type de beton", "B15", "B20", "B25", "B30", "B40","B45","B50","B55","B60","B65","B70","B75","B80","B85","B90", "B95","B100"],
            height=40,
            #  command=self.bti(file,), 
             command=self.bti,
             fg_color="#293DA4"
             )
        self.bt_option.grid(row=0, column=0, padx=20, pady=(10, 10))
        
        self.entry = customtkinter.CTkEntry(
            master=self.head, 
            placeholder_text="Enter le type de beton", 
            height=40, width=400,
            font=customtkinter.CTkFont(size=18, weight="bold")
            )
        self.entry.grid(row=0, column=1, columnspan=2, padx=(20, 0), pady=(20, 20), sticky="ew")

        self.main_button_1 = customtkinter.CTkButton(
            master=self.head, 
            fg_color="transparent", 
            border_width=2, 
            text_color="#293DA4",
            border_color="#293DA4" ,
            height=40,
            text="Search",
            font=customtkinter.CTkFont(size=18, weight="bold")
            )
        self.main_button_1.grid(row=0, column=3, padx=(20, 20), pady=(20, 20), sticky="nsew")
        
                ########################################################
        #configuration
        self.bar_config = customtkinter.CTkFrame(master=self, height=40, border_color="#DC3F3F", border_width=1)
        self.bar_config.grid(row=3, column=1, columnspan=2, padx=(20, 0), pady=(0, 0), sticky="nsew")
        
                ##################################################
        #bouton pour valider la recher
        self.btn_recherche = customtkinter.CTkOptionMenu(
            master=self.bar_config,
            values=[
                "Exporter", "Exel", "PDF", "Word"
            ],
            fg_color="#293DA4",
            # command=lambda:self.etat(),
            command=lambda eta : self.etat(eta)
            # font=customtkinter.CTkFont(size=18, weight="bold")
        ).grid(row=0, column=1, padx=(20, 20), pady=(10, 10), sticky="nsew")
        
        self.save = customtkinter.CTkButton(master=self.bar_config, text="Sauvegarder", command=self.saves)
        self.save.grid(row=0, column=2)
        
        # #         # create scrollable frame
        self.scrollable_frame = customtkinter.CTkScrollableFrame(
            self, 
            label_text="Caracteristique",
            label_font=customtkinter.CTkFont(size=24, weight="bold"),
            label_text_color="#DC3F3F"
            )
        self.scrollable_frame.grid(row=1, column=1, padx=(20, 0), pady=(20, 0), sticky="nsew")
        self.scrollable_frame.grid_columnconfigure(1, weight=1)
        self.scrollable_frame.grid_columnconfigure((1, 3), weight=0)
        self.scrollable_frame.grid_rowconfigure((0, 0, 0), weight=1)
        # self.scrollable_frame.grid_columnconfigure(0, weight=1)
        # create textbox
        self.textbox = customtkinter.CTkFrame(master=self.scrollable_frame, width=250, fg_color="#DC3F3F")
        self.textbox.grid(row=0, column=1, padx=(0, 0), pady=(20, 0), sticky="nsew")
        self.textbox.grid_columnconfigure(1, weight=1)
        self.textbox.grid_columnconfigure((1, 3), weight=0)
        self.textbox.grid_rowconfigure((0, 0, 0), weight=1)
        
        ###############################
        #frame pour les l'onglet cell name
        self.cell_name = customtkinter.CTkFrame(master=self.textbox,  border_color="red",corner_radius=2 )
        self.cell_name.grid(row=0, column=0, sticky="nsew", pady=(5, 5))
        
        
        #label
        self.cell_name_on = customtkinter.CTkFrame(master=self.cell_name,  border_width=1)
        self.cell_name_on.grid(row=0, column=0, sticky="nsew", padx=(10, 10),)
        self.cell_name_lb = customtkinter.CTkLabel(master=self.cell_name_on, text="Dosage en ciment(KG/m3)",  height=50,font=customtkinter.CTkFont(size=14, weight="bold"))
        self.cell_name_lb.grid(row=0, column=0, sticky="nsew")
        #
        self.cell_name_on = customtkinter.CTkFrame(master=self.cell_name_on,  border_width=1, height=100,)
        self.cell_name_on.grid(row=1, column=0, sticky="nsew")
        self.cim = tk.Label(master=self.cell_name_on, text="", font=customtkinter.CTkFont(size=20, weight="bold"))
        self.cim.grid(row=0, column=0, sticky="nsew", pady=(30, 10))
        
        
        
        #frame pour les l'onglet site name
        self.site_name = customtkinter.CTkFrame(master=self.textbox,  border_color="red", )
        self.site_name.grid(row=0, column=1, sticky="nsew",pady=(5, 5), )
        
        #label
        self.site_name_on = customtkinter.CTkFrame(master=self.site_name,  border_width=1, height=50)
        self.site_name_on.grid(row=0, column=0, padx=(10, 10))
        self.cell_name_lb = customtkinter.CTkLabel(master=self.site_name_on, text="Dosage Gravier 00/15(KG/m3)",height=50,font=customtkinter.CTkFont(size=14, weight="bold"))
        self.cell_name_lb.grid(row=0, column=0, sticky="nsew")
        
        #
        self.cell_name_on = customtkinter.CTkFrame(master=self.site_name_on,  border_width=1, height=100)
        self.cell_name_on.grid(row=1, column=0, sticky="nsew")
        self.gra = tk.Label(master=self.cell_name_on, text="", font=customtkinter.CTkFont(size=20, weight="bold"))
        self.gra.grid(row=0, column=0, sticky="nsew", pady=(30, 10))
        
        #frame pour les l'onglet SR
        self.sr = customtkinter.CTkFrame(master=self.textbox,  border_color="red", )
        self.sr.grid(row=0, column=2, sticky="nsew",pady=(5, 5), )
        
        #label
        self.sr_on = customtkinter.CTkFrame(master=self.sr,  border_width=1, height=50)
        self.sr_on.grid(row=0, column=0, padx=(10, 10))
        self.cell_name_lb = customtkinter.CTkLabel(master=self.sr_on, text="Dosage Gravier 25/15(KG/m3)",  height=50,font=customtkinter.CTkFont(size=14, weight="bold"))
        self.cell_name_lb.grid(row=0, column=0, sticky="nsew")
        
        #
        self.cell_name_on = customtkinter.CTkFrame(master=self.sr_on,  border_width=1, height=100)
        self.cell_name_on.grid(row=1, column=0, sticky="nsew")
        self.gra2 = tk.Label(master=self.cell_name_on, text="", font=customtkinter.CTkFont(size=20, weight="bold"))
        self.gra2.grid(row=0, column=0, sticky="nsew", pady=(30, 10))
        
        #frame pour les l'onglet Congestion rate(sdcch)
        self.cr_sdcch = customtkinter.CTkFrame(master=self.textbox,  border_color="red", )
        self.cr_sdcch.grid(row=0, column=3, sticky="nsew",pady=(5, 5), )
        
        #label
        self.cr_on = customtkinter.CTkFrame(master=self.cr_sdcch,  border_width=1, height=50)
        self.cr_on.grid(row=0, column=0, padx=(10, 10))
        self.cell_name_lb = customtkinter.CTkLabel(master=self.cr_on, text="Dosage en Eau(L)",  height=50,font=customtkinter.CTkFont(size=14, weight="bold"))
        self.cell_name_lb.grid(row=0, column=0, sticky="nsew")
        
        #
        self.cell_name_on = customtkinter.CTkFrame(master=self.cr_on,  border_width=1, height=100)
        self.cell_name_on.grid(row=1, column=0, sticky="nsew")
        self.eau = tk.Label(master=self.cell_name_on, text="", font=customtkinter.CTkFont(size=20, weight="bold"))
        self.eau.grid(row=0, column=0, sticky="nsew", pady=(30, 10))
        
        #frame pour les l'onglet Congestion rate(tch)
        self.cr_tch = customtkinter.CTkFrame(master=self.textbox, border_color="red", )
        self.cr_tch.grid(row=0, column=4, sticky="nsew",pady=(5, 5), )
        
        #label
        self.cr_thc_on = customtkinter.CTkFrame(master=self.cr_tch,  border_width=1, height=50)
        self.cr_thc_on.grid(row=0, column=0, padx=(10, 10))
        self.cell_name_lb = customtkinter.CTkLabel(master=self.cr_thc_on, text="Dosage Sable(KG/m3)",  height=50,font=customtkinter.CTkFont(size=14, weight="bold"))
        self.cell_name_lb.grid(row=0, column=0, sticky="nsew")
        
        #
        self.cell_name_on = customtkinter.CTkFrame(master=self.cr_thc_on,  border_width=1, height=100)
        self.cell_name_on.grid(row=1, column=0, sticky="nsew")
        self.sable = tk.Label(master=self.cell_name_on, text="", font=customtkinter.CTkFont(size=20, weight="bold"))
        self.sable.grid(row=0, column=0, sticky="nsew", pady=(30, 10))
        
        #frame pour les l'onglet availebility
        self.avail = customtkinter.CTkFrame(master=self.textbox, border_color="red",)
        self.avail.grid(row=0, column=5, sticky="nsew",pady=(5, 5), )
        
        #label
        self.avail_on = customtkinter.CTkFrame(master=self.avail,  border_width=1, height=50, )
        self.avail_on.grid(row=0, column=0, padx=(10, 10))
        self.cell_name_lb = customtkinter.CTkLabel(master=self.avail_on, text="Type de Beton", height=50,font=customtkinter.CTkFont(size=14, weight="bold"))
        self.cell_name_lb.grid(row=0, column=0, sticky="nsew")
        
        #
        self.cell_name_on = customtkinter.CTkFrame(master=self.avail_on,  border_width=1, height=100)
        self.cell_name_on.grid(row=1, column=0, sticky="nsew")
        self.type = tk.Label(master=self.cell_name_on, text="", font=customtkinter.CTkFont(size=20, weight="bold"))
        self.type.grid(row=0, column=0, sticky="nsew", pady=(30, 10))
        
        
        
        ###########################################################################################################
        self.conf = customtkinter.CTkFrame(master=self.scrollable_frame, )
        self.conf.grid(row=1, column=1, padx=(10, 0), pady=(50, 0), sticky="nsew")
        self.conf.grid_columnconfigure(1, weight=1)
        self.conf.grid_columnconfigure((1, 2), weight=0)
        self.conf.grid_rowconfigure((0, 0, 0), weight=1)
        
        #################
        self.main = customtkinter.CTkScrollableFrame(
            master=self.conf,
            label_text="Main d'oeuvre",
            width=250,
            label_font=customtkinter.CTkFont(size=18, weight="bold"),
            label_text_color="#DC3F3F"
            )
                        ###################
                        
                        
        self.main.grid(row=0, column=0, padx=20,)
        self.checkbox_1 = customtkinter.CTkCheckBox(master=self.main, text="Conducteur de traveau")
        self.checkbox_1.grid(row=1, column=0, pady=(20, 0), padx=20, sticky="nsew",)
        self.checkbox_2 = customtkinter.CTkCheckBox(master=self.main, text="Ingenieur de suivie")
        self.checkbox_2.grid(row=2, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_3 = customtkinter.CTkCheckBox(master=self.main, text="Superviseur QHSE")
        self.checkbox_3.grid(row=3, column=0, pady=20, padx=20, sticky="nsew")
        self.checkbox_4 = customtkinter.CTkCheckBox(master=self.main, text="Chef d'equipe")
        self.checkbox_4.grid(row=4, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_5 = customtkinter.CTkCheckBox(master=self.main, text="Mecaniciens")
        self.checkbox_5.grid(row=5, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_6 = customtkinter.CTkCheckBox(master=self.main, text="Electriciens")
        self.checkbox_6.grid(row=6, column=0, pady=20, padx=20, sticky="nsew")
        self.checkbox_7 = customtkinter.CTkCheckBox(master=self.main, text="Plombiers")
        self.checkbox_7.grid(row=7, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_8 = customtkinter.CTkCheckBox(master=self.main, text="Conducteur d'engain")
        self.checkbox_8.grid(row=8, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_9 = customtkinter.CTkCheckBox(master=self.main, text="Chauffeur camion")
        self.checkbox_9.grid(row=9, column=0, pady=20, padx=20, sticky="nsew")
        self.checkbox_10 = customtkinter.CTkCheckBox(master=self.main, text="Chef chantier")
        self.checkbox_10.grid(row=10, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_11 = customtkinter.CTkCheckBox(master=self.main, text="Manoeuvre")
        self.checkbox_11.grid(row=11, column=0, pady=(20, 0), padx=20, sticky="nsew")
                #################
        
        ######################################
        
        
        
        self.checkbox_12 = customtkinter.CTkEntry(master=self.main, width=50)
        self.checkbox_12.grid(row=1, column=1, pady=(20, 0), padx=0, sticky="wn",)
        self.checkbox_13 = customtkinter.CTkEntry(master=self.main, width=50)
        self.checkbox_13.grid(row=2, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_14 = customtkinter.CTkEntry(master=self.main, width=50)
        self.checkbox_14.grid(row=3, column=1, pady=20, padx=0, sticky="wn")
        self.checkbox_15 = customtkinter.CTkEntry(master=self.main, width=50)
        self.checkbox_15.grid(row=4, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_16 = customtkinter.CTkEntry(master=self.main, width=50)
        self.checkbox_16.grid(row=5, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_17 = customtkinter.CTkEntry(master=self.main, width=50)
        self.checkbox_17.grid(row=6, column=1, pady=20, padx=0, sticky="wn")
        self.checkbox_18 = customtkinter.CTkEntry(master=self.main, width=50)
        self.checkbox_18.grid(row=7, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_19 = customtkinter.CTkEntry(master=self.main, width=50)
        self.checkbox_19.grid(row=8, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_20 = customtkinter.CTkEntry(master=self.main, width=50)
        self.checkbox_20.grid(row=9, column=1, pady=20, padx=0, sticky="wn")
        self.checkbox_21 = customtkinter.CTkEntry(master=self.main, width=50)
        self.checkbox_21.grid(row=10, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_22 = customtkinter.CTkEntry(master=self.main, width=50)
        self.checkbox_22.grid(row=11, column=1, pady=(20, 0), padx=0, sticky="wn")

        #########################################
        
        ###############
        self.engin = customtkinter.CTkScrollableFrame(
            master=self.conf,
            label_text="Engains",
            width=230,
            label_font=customtkinter.CTkFont(size=18, weight="bold"),
            label_text_color="#DC3F3F"
            )
        
        self.engin.grid(row=0, column=1, padx=20,)
        self.checkbox_23 = customtkinter.CTkCheckBox(master=self.engin, text="Tractopelle")
        self.checkbox_23.grid(row=1, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_24 = customtkinter.CTkCheckBox(master=self.engin, text="Pelle excavatrice")
        self.checkbox_24.grid(row=2, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_25 = customtkinter.CTkCheckBox(master=self.engin, text="Betonniere")
        self.checkbox_25.grid(row=3, column=0, pady=20, padx=20, sticky="nsew")
        self.checkbox_26 = customtkinter.CTkCheckBox(master=self.engin, text="Camion toupie")
        self.checkbox_26.grid(row=4, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_27 = customtkinter.CTkCheckBox(master=self.engin, text="Pelle chargeuse")
        self.checkbox_27.grid(row=5, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_28 = customtkinter.CTkCheckBox(master=self.engin, text="Nivelleuse")
        self.checkbox_28.grid(row=6, column=0, pady=20, padx=20, sticky="nsew")
        self.checkbox_29 = customtkinter.CTkCheckBox(master=self.engin, text="Compacteurs")
        self.checkbox_29.grid(row=7, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_30 = customtkinter.CTkCheckBox(master=self.engin, text="Autobetonnier")
        self.checkbox_30.grid(row=8, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_31 = customtkinter.CTkCheckBox(master=self.engin, text="Central a beton")
        self.checkbox_31.grid(row=9, column=0, pady=20, padx=20, sticky="nsew")
        self.checkbox_32 = customtkinter.CTkCheckBox(master=self.engin, text="Bouette")
        self.checkbox_32.grid(row=10, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_33 = customtkinter.CTkCheckBox(master=self.engin, text="Pelles")
        self.checkbox_33.grid(row=11, column=0, pady=(20, 0), padx=20, sticky="nsew")
        
        
        ###############################################
        self.checkbox_34 = customtkinter.CTkEntry(master=self.engin, width=50)
        self.checkbox_34.grid(row=1, column=1, pady=(20, 0), padx=0, sticky="wn",)
        self.checkbox_35 = customtkinter.CTkEntry(master=self.engin, width=50)
        self.checkbox_35.grid(row=2, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_36 = customtkinter.CTkEntry(master=self.engin, width=50)
        self.checkbox_36.grid(row=3, column=1, pady=20, padx=0, sticky="wn")
        self.checkbox_37 = customtkinter.CTkEntry(master=self.engin, width=50)
        self.checkbox_37.grid(row=4, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_38 = customtkinter.CTkEntry(master=self.engin, width=50)
        self.checkbox_38.grid(row=5, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_39 = customtkinter.CTkEntry(master=self.engin, width=50)
        self.checkbox_39.grid(row=6, column=1, pady=20, padx=0, sticky="wn")
        self.checkbox_40 = customtkinter.CTkEntry(master=self.engin, width=50)
        self.checkbox_40.grid(row=7, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_42 = customtkinter.CTkEntry(master=self.engin, width=50)
        self.checkbox_42.grid(row=8, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_43 = customtkinter.CTkEntry(master=self.engin, width=50)
        self.checkbox_43.grid(row=9, column=1, pady=20, padx=0, sticky="wn")
        self.checkbox_44 = customtkinter.CTkEntry(master=self.engin, width=50)
        self.checkbox_44.grid(row=10, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_45 = customtkinter.CTkEntry(master=self.engin, width=50)
        self.checkbox_45.grid(row=11, column=1, pady=(20, 0), padx=0, sticky="wn")
        ##############################################
        
        
        #################
        self.materiau = customtkinter.CTkScrollableFrame(
            master=self.conf,
            label_text="Materiaux",
            # width=250
            label_font=customtkinter.CTkFont(size=18, weight="bold"),
            label_text_color="#DC3F3F"
            )
        
        self.materiau.grid(row=0, column=2, padx=20,)
        self.checkbox_46 = customtkinter.CTkCheckBox(master=self.materiau, text="Sable")
        self.checkbox_46.grid(row=1, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_47 = customtkinter.CTkCheckBox(master=self.materiau, text="Gravier 00/15")
        self.checkbox_47.grid(row=2, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_48 = customtkinter.CTkCheckBox(master=self.materiau, text="Gravier 15/25")
        self.checkbox_48.grid(row=3, column=0, pady=20, padx=20, sticky="nsew")
        self.checkbox_49 = customtkinter.CTkCheckBox(master=self.materiau, text="Ciment")
        self.checkbox_49.grid(row=4, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_50 = customtkinter.CTkCheckBox(master=self.materiau, text="Eau")
        self.checkbox_50.grid(row=5, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_52 = customtkinter.CTkCheckBox(master=self.materiau, text=" Fer de 6")
        self.checkbox_52.grid(row=6, column=0, pady=20, padx=20, sticky="nsew")
        self.checkbox_53 = customtkinter.CTkCheckBox(master=self.materiau, text="Fer de 8")
        self.checkbox_53.grid(row=7, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_54 = customtkinter.CTkCheckBox(master=self.materiau, text="Fer de 10")
        self.checkbox_54.grid(row=8, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_55 = customtkinter.CTkCheckBox(master=self.materiau, text="Fer de 12")
        self.checkbox_55.grid(row=9, column=0, pady=20, padx=20, sticky="nsew")
        self.checkbox_56 = customtkinter.CTkCheckBox(master=self.materiau, text="Fer de 14")
        self.checkbox_56.grid(row=10, column=0, pady=(20, 0), padx=20, sticky="nsew")
        self.checkbox_57 = customtkinter.CTkCheckBox(master=self.materiau, text="Fer de 16")
        self.checkbox_57.grid(row=11, column=0, pady=(20, 0), padx=20, sticky="nsew")
        
                ###############################################
        self.checkbox_58= customtkinter.CTkEntry(master=self.materiau, width=50, placeholder_text="En KG")
        self.checkbox_58.grid(row=1, column=1, pady=(20, 0), padx=0, sticky="wn",)
        self.checkbox_59 = customtkinter.CTkEntry(master=self.materiau, width=50, placeholder_text="En KG")
        self.checkbox_59.grid(row=2, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_60 = customtkinter.CTkEntry(master=self.materiau, width=50, placeholder_text="En KG")
        self.checkbox_60.grid(row=3, column=1, pady=20, padx=0, sticky="wn")
        self.checkbox_61 = customtkinter.CTkEntry(master=self.materiau, width=50, placeholder_text="En KG")
        self.checkbox_61.grid(row=4, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_62 = customtkinter.CTkEntry(master=self.materiau, width=50, placeholder_text="En M3")
        self.checkbox_62.grid(row=5, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_63 = customtkinter.CTkEntry(master=self.materiau, width=50)
        self.checkbox_63.grid(row=6, column=1, pady=20, padx=0, sticky="wn")
        self.checkbox_64 = customtkinter.CTkEntry(master=self.materiau, width=50)
        self.checkbox_64.grid(row=7, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_65 = customtkinter.CTkEntry(master=self.materiau, width=50)
        self.checkbox_65.grid(row=8, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_66 = customtkinter.CTkEntry(master=self.materiau, width=50)
        self.checkbox_66.grid(row=9, column=1, pady=20, padx=0, sticky="wn")
        self.checkbox_67 = customtkinter.CTkEntry(master=self.materiau, width=50)
        self.checkbox_67.grid(row=10, column=1, pady=(20, 0), padx=0, sticky="wn")
        self.checkbox_68 = customtkinter.CTkEntry(master=self.materiau, width=50)
        self.checkbox_68.grid(row=11, column=1, pady=(20, 0), padx=0, sticky="wn")
        ##############################################
        
        ##################
        #bouton de confirmation
        self.btn = customtkinter.CTkButton(master=self.conf, text="Consulter",fg_color="#293DA4",command=self.valid, font=customtkinter.CTkFont(size=14, weight="bold"))
        self.btn.grid(row=0, column=3, sticky="s")





    mat = None
    main = None
    eng = None
    def valid(self):
            try: 
                new = customtkinter.CTkToplevel(self)
                # new.title("Erreur de fichier")
                new.geometry("1080x720")
                new.resizable(False, False)
                new.attributes('-topmost', True)
                new.grab_set()
                Conducteur = 0
                if self.checkbox_1.get() == 1:
                    Conducteur = self.checkbox_12.get() 
                    Conducteur = int(Conducteur)*10000
                    
                Ingenieur = 0
                if self.checkbox_2.get() == 1:
                    Ingenieur = self.checkbox_13.get()
                    Ingenieur = int(Ingenieur)*10000
                    
                Superviseur = 0
                if self.checkbox_3.get() == 1:
                    Superviseur = self.checkbox_14.get()
                    Superviseur = int(Superviseur)*10000
                    
                Chef_equipe = 0
                if self.checkbox_4.get() == 1:
                    Chef_equipe = self.checkbox_15.get()
                    Chef_equipe  = int(Chef_equipe)*10000
                    
                Mecaniciens = 0
                if self.checkbox_5.get() == 1:
                    Mecaniciens = self.checkbox_16.get()
                    Mecaniciens = int(Mecaniciens)*5000
                    
                Electriciens = 0
                if self.checkbox_6.get() == 1:
                    Electriciens = self.checkbox_17.get()
                    Electriciens = int(Electriciens)*5000
                    
                Plombiers = 0
                if self.checkbox_7.get() == 1:
                    Plombiers = self.checkbox_18.get()
                    Plombiers = int(Plombiers)*5000
                    
                Conducteur1 = 0
                if self.checkbox_8.get() == 1:
                    Conducteur1 = self.checkbox_19.get()
                    Conducteur1 = int(Conducteur1)*10000
                    
                Chauffeur =  0
                if self.checkbox_9.get() == 1:
                    Chauffeur = self.checkbox_20.get()
                    Chauffeur = int(Chauffeur)*10000
                    
                Chef_chantier = 0
                if self.checkbox_10.get() == 1:
                    Chef_chantier = self.checkbox_21.get()
                    Chef_chantier = int(Chef_chantier)*10000
                    
                Manoeuvre = 0
                if self.checkbox_11.get() == 1:
                    Manoeuvre = self.checkbox_22.get()
                    Manoeuvre = int(Manoeuvre)*3000
                quant1 = [
                    self.checkbox_12.get(),
                    self.checkbox_13.get(),
                    self.checkbox_14.get(),
                    self.checkbox_15.get(),
                    self.checkbox_16.get(),
                    self.checkbox_17.get(),
                    self.checkbox_18.get(),
                    self.checkbox_19.get(),
                    self.checkbox_20.get(),
                    self.checkbox_21.get(),
                    self.checkbox_22.get()
                ]
                    
                Tractopelle =0
                Pelle_excavatrice =0
                Betonniere=0
                Camion_toupie=0
                Pelle_chargeuse=0
                Nivelleuse=0
                Compacteurs=0
                Autobetonnier=0
                Central_a_beton=0
                Bouette=0
                Pelles=0
                if self.checkbox_23.get() == 1:
                    Tractopelle = self.checkbox_34.get()
                    Tractopelle = int(Tractopelle)*200000
                
                if self.checkbox_24.get() == 1:
                    Pelle_excavatrice = self.checkbox_35.get()
                    Pelle_excavatrice = int(Pelle_excavatrice)*360000
                    
                if self.checkbox_25.get() == 1:
                    Betonniere =  self.checkbox_36.get()
                    Betonniere = int(Betonniere)*300000
                    
                if self.checkbox_26.get() == 1:
                    Camion_toupie = self.checkbox_37.get()
                    Camion_toupie = int(Camion_toupie)*150000
                    
                if self.checkbox_27.get() == 1:
                    Pelle_chargeuse = self.checkbox_38.get()
                    Pelle_chargeuse = int(Pelle_chargeuse)*300000
                    
                if self.checkbox_27.get() == 1:
                    Nivelleuse = self.checkbox_39.get()
                    Nivelleuse = int(Nivelleuse)*400000
                    
                if self.checkbox_29.get() == 1:
                    Compacteurs = self.checkbox_40.get()
                    Compacteurs = int(Compacteurs)*250000
                
                if self.checkbox_30.get() == 1:
                    Autobetonnier = self.checkbox_42.get()
                    Autobetonnier = int(Autobetonnier)*100000
                    
                if self.checkbox_31.get() == 1:
                    Central_a_beton = self.checkbox_43.get()
                    Central_a_beton = int(Central_a_beton)*1300000
                    
                if self.checkbox_32.get() == 1:
                    Bouette = self.checkbox_44.get()
                    Bouette = int(Bouette)*30000
                    
                    
                if self.checkbox_33.get() == 1:
                    Pelles = self.checkbox_45.get()
                    Pelles = int(Pelles)*5000
                    
                quant2 = [
                    self.checkbox_34.get(),
                    self.checkbox_35.get(),
                    self.checkbox_36.get(),
                    self.checkbox_37.get(),
                    self.checkbox_38.get(),
                    self.checkbox_39.get(),
                    self.checkbox_40.get(),
                    self.checkbox_42.get(),
                    self.checkbox_43.get(),
                    self.checkbox_44.get(),
                    self.checkbox_45.get()
                ]
                    
                    ##################################
                Sable = 0
                Gravier_0015 = 0
                Gravier_1525 = 0
                Ciment=0
                Eau=0
                Fer_de_6 =0
                Fer_de_8 =0
                Fer_de_10 =0
                Fer_de_12 =0
                Fer_de_14 =0
                Fer_de_16 = 0
                if self.checkbox_46.get() == 1:
                    Sable = self.checkbox_58.get()
                    Sable = int(Sable)*275
                    
                if self.checkbox_47.get() == 1:
                    Gravier_0015 = self.checkbox_59.get()
                    Gravier_0015 = int(Gravier_0015)*45
                    
                if self.checkbox_48.get() == 1:
                    Gravier_1525 = self.checkbox_60.get()
                    Gravier_1525 = int(Gravier_1525)*305
                    
                if self.checkbox_49.get() == 1:
                    Ciment = self.checkbox_61.get()
                    Ciment = int(Ciment)*5500
                    
                if self.checkbox_50.get() == 1:
                    Eau = self.checkbox_62.get()
                    Eau = int(Eau)*830
                    
                    
                if self.checkbox_52.get() == 1:
                    Fer_de_6 = self.checkbox_63.get()
                    Fer_de_6 = int(Fer_de_6)*600
                    
                if self.checkbox_53.get() == 1:
                    Fer_de_8 = self.checkbox_64.get()
                    Fer_de_8 = int(Fer_de_8)*800
                    
                if self.checkbox_54.get() == 1:
                    Fer_de_10 = self.checkbox_65.get()
                    Fer_de_10 = int(Fer_de_10)*1000
                
                if self.checkbox_55.get() == 1:
                    Fer_de_12 = self.checkbox_66.get()
                    Fer_de_12 = int(Fer_de_12)*1200
                    
                if self.checkbox_56.get() == 1:
                    Fer_de_14 = self.checkbox_67.get()
                    Fer_de_14 = int(Fer_de_14)*1400
                    
                if self.checkbox_57.get() == 1:
                    Fer_de_16 = self.checkbox_68.get()
                    Fer_de_16 = int(Fer_de_16)*1600
                    
                quant3 = [
                    self.checkbox_58.get(),
                    self.checkbox_59.get(),
                    self.checkbox_60.get(),
                    self.checkbox_61.get(),
                    self.checkbox_62.get(),
                    self.checkbox_63.get(),
                    self.checkbox_64.get(),
                    self.checkbox_65.get(),
                    self.checkbox_66.get(),
                    self.checkbox_67.get(),
                    self.checkbox_68.get()
                ]
                global mat
                global eng
                global main
                mat = [
                    {
                        "Materiaux":"Sable",
                        "Prix":Sable,
                        "Quantite":self.checkbox_58.get()
                    },
                    {
                        "Materiaux":"Gravier_0015",
                        "Prix":Gravier_0015,
                        "Quantite":self.checkbox_59.get()
                    },
                    {
                        "Materiaux":"Gravier_1525",
                        "Prix":Gravier_1525,
                        "Quantite":self.checkbox_60.get()
                    },
                    {
                        "Materiaux":"Ciment",
                        "Prix":Ciment,
                        "Quantite":self.checkbox_61.get()
                    },
                    {
                        "Materiaux":"Eau",
                        "Prix":Eau,
                        "Quantite":self.checkbox_62.get()
                    },
                    {
                        "Materiaux":"Fer_de_6",
                        "Prix":Fer_de_6,
                        "Quantite":self.checkbox_63.get()
                    },
                    {
                        "Materiaux":"Fer_de_8",
                        "Prix":Fer_de_8,
                        "Quantite":self.checkbox_64.get()
                    },
                    {
                        "Materiaux":"Fer_de_10",
                        "Prix":Fer_de_10,
                        "Quantite":self.checkbox_65.get()
                    },
                    {
                        "Materiaux":"Fer_de_12",
                        "Prix":Fer_de_12,
                        "Quantite":self.checkbox_66.get()
                    },
                    {
                        "Materiaux":"Fer_de_14",
                        "Prix":Fer_de_14,
                        "Quantite":self.checkbox_67.get()
                    },
                    {
                        "Materiaux":"Fer_de_16",
                        "Prix":Fer_de_16,
                        "Quantite":self.checkbox_68.get()
                    },
                ]
                    
                    
                    
                materiau = {
                    "Sable" :Sable,
                    "Gravier_0015" :Gravier_0015,
                    "Gravier_1525" :Gravier_1525,
                    "Ciment":Ciment,
                    "Eau":Eau,
                    "Fer_de_6" :Fer_de_6,
                    "Fer_de_8" :Fer_de_8,
                    "Fer_de_10" :Fer_de_10,
                    "Fer_de_12" :Fer_de_12,
                    "Fer_de_14" :Fer_de_14,
                    "Fer_de_16" :Fer_de_16,
                }
                
                main = [
                    {
                        "main_oeuvre":"Conducteur",
                        "prix":Conducteur,
                        "quntite":self.checkbox_12.get(),
                    },
                    {
                        "main_oeuvre":"Ingenieur",
                        "prix":Ingenieur,
                        "quntite":self.checkbox_13.get(),
                    },
                    {
                        "main_oeuvre":"Superviseur",
                        "prix":Superviseur,
                        "quntite":self.checkbox_14.get(),
                    },
                    {
                        "main_oeuvre":"Chef_equipe",
                        "prix":Chef_equipe,
                        "quntite":self.checkbox_15.get(),
                    },
                    {
                        "main_oeuvre":"Mecaniciens",
                        "prix":Mecaniciens,
                        "quntite":self.checkbox_16.get(),
                    },
                    {
                        "main_oeuvre":"Electriciens",
                        "prix":Electriciens,
                        "quntite":self.checkbox_17.get(),
                    },
                    {
                        "main_oeuvre":"Plombiers",
                        "prix":Plombiers,
                        "quntite":self.checkbox_18.get(),
                    },
                    {
                        "main_oeuvre":"Conducteur1",
                        "prix":Conducteur1,
                        "quntite":self.checkbox_19.get(),
                    },
                    {
                        "main_oeuvre":"Chauffeur",
                        "prix":Chauffeur,
                        "quntite":self.checkbox_20.get(),
                    },
                    {
                        "main_oeuvre":"Chef_chantier",
                        "prix":Chef_chantier,
                        "quntite":self.checkbox_21.get(),
                    },
                    {
                        "main_oeuvre":"Manoeuvre",
                        "prix":Manoeuvre,
                        "quntite":self.checkbox_22.get(),
                    },
                ]
                
                
                main_oeuvre = {
                    "Conducteur" :Conducteur,
                    "Ingenieur":Ingenieur,
                    "Superviseur":Superviseur,
                    "Chef_equipe":Chef_equipe,
                    "Mecaniciens":Mecaniciens, 
                    "Electriciens":Electriciens,
                    "Plombiers":Plombiers,
                    "Conducteur1":Conducteur1,
                    "Chauffeur":Chauffeur,
                    "Chef_chantier":Chef_chantier,
                    "Manoeuvre":Manoeuvre
                    } 
                
                eng = [
                    {
                        "Engins":"Tractopelle",
                        "Prix":Tractopelle,
                        "Quantite": self.checkbox_58.get(),
                    },
                     {
                        "Engins":"Pelle_excavatrice",
                        "Prix":Pelle_excavatrice,
                        "Quantite": self.checkbox_59.get(),
                    },
                      {
                        "Engins":"Betonniere",
                        "Prix":Betonniere,
                        "Quantite": self.checkbox_60.get(),
                    },
                       {
                        "Engins":"Camion_toupie",
                        "Prix":Camion_toupie,
                        "Quantite": self.checkbox_61.get(),
                    },
                        {
                        "Engins":"Pelle_chargeuse",
                        "Prix":Pelle_chargeuse,
                        "Quantite": self.checkbox_62.get(),
                    },
                         {
                        "Engins":"Nivelleuse",
                        "Prix":Nivelleuse,
                        "Quantite": self.checkbox_63.get(),
                    },
                          {
                        "Engins":"Compacteurs",
                        "Prix":Compacteurs,
                        "Quantite": self.checkbox_64.get(),
                    },
                           {
                        "Engins":"Autobetonnier",
                        "Prix":Autobetonnier,
                        "Quantite": self.checkbox_65.get(),
                    },
                            {
                        "Engins":"Central_a_beton",
                        "Prix":Central_a_beton,
                        "Quantite": self.checkbox_66.get(),
                    },
                             {
                        "Engins":"Bouette",
                        "Prix":Bouette,
                        "Quantite": self.checkbox_67.get(),
                    },
                              {
                        "Engins":"Pelles",
                        "Prix":Pelles,
                        "Quantite": self.checkbox_68.get(),
                    },
                ]
                
                
                engins = {
                    "Tractopelle" :Tractopelle,
                    "Pelle_excavatrice" :Pelle_excavatrice,
                    "Betonniere":Betonniere,
                    "Camion_toupie":Camion_toupie,
                    "Pelle_chargeuse":Pelle_chargeuse,
                    "Nivelleuse":Nivelleuse,
                    "Compacteurs":Compacteurs,
                    "Autobetonnier":Autobetonnier,
                    "Central_a_beton":Central_a_beton,
                    "Bouette":Bouette,
                    "Pelles":Pelles,
                }
                
                
                new.grid_columnconfigure(1, weight=1)
                new.grid_columnconfigure((2, 3), weight=0)
                new.grid_rowconfigure((0, 1, 1), weight=1)
        
                textbox = customtkinter.CTkFrame(new, width=250)
                textbox.grid(row=0, column=1, padx=(20, 0), pady=(20, 0), sticky="nsew")
                textbox.grid_columnconfigure(1, weight=1)
                textbox.grid_columnconfigure((1, 2), weight=0)
                textbox.grid_rowconfigure((0, 0, 0), weight=1)
        
        ###############################
        #frame pour les l'onglet main d'oeuvre
                cell_name = customtkinter.CTkFrame(master=textbox,  border_color="red",corner_radius=2 )
                cell_name.grid(row=0, column=0, sticky="nsew", pady=(5, 5))
        
        
        #label
                cell_name_on = customtkinter.CTkFrame(master=cell_name,  border_width=1,  height=50)
                cell_name_on.grid(row=0, column=0, sticky="nsew")
                cell_name_lb = customtkinter.CTkLabel(master=cell_name_on, text="Main d'oeuvre",width=350,  height=50,font=customtkinter.CTkFont(size=14, weight="bold"))
                cell_name_lb.grid(row=0, column=0, sticky="nsew",)
            #
                cell_name_on = customtkinter.CTkFrame(master=cell_name_on,  border_width=1, border_color="#DC3F3F")
                cell_name_on.grid(row=1, column=0, sticky="nsew",padx=(5, 5), pady=(5, 5))


                total1 = customtkinter.CTkLabel(master=cell_name, text=f"Total  :  {sum(main_oeuvre.values())} XAF")
                total1.grid(row=1, column=0)
                ligne = 0
                for k, v in main_oeuvre.items():
                    vue = customtkinter.CTkLabel(master=cell_name_on, text=f"{k}: ",font=customtkinter.CTkFont(size=14) )
                    vue.grid(row=ligne, column=0, padx=(10, 5), pady=(10,10))
                    label = customtkinter.CTkLabel(master=cell_name_on, text=f"{v}", font=customtkinter.CTkFont(size=14))
                    label.grid(row=ligne, column=2,padx=(10, 5))
                    quant = customtkinter.CTkLabel(master=cell_name_on, text=f"{quant1[ligne]}", font=customtkinter.CTkFont(size=14))
                    quant.grid(row=ligne, column=3,padx=(50, 5), sticky="e")
                    ligne += 1
                        
        
        
        #frame pour les l'onglet Engins
                site_name = customtkinter.CTkFrame(master=textbox,  border_color="red",)
                site_name.grid(row=0, column=1, sticky="nsew",pady=(5, 5), )
        
        #label
                site_name_on = customtkinter.CTkFrame(master=site_name,  border_width=1, height=50)
                site_name_on.grid(row=0, column=0)
                cell_name_lb = customtkinter.CTkLabel(master=site_name_on, text="Engins",width=350, height=50,font=customtkinter.CTkFont(size=14, weight="bold"))
                cell_name_lb.grid(row=0, column=0, sticky="nsew")
        
        #
                minor_on = customtkinter.CTkFrame(master=site_name_on,  border_width=1, border_color="#DC3F3F")
                minor_on.grid(row=1, column=0, sticky="nsew",padx=(5, 5), pady=(5, 5))
                
                total2 = customtkinter.CTkLabel(master=site_name, text=f"Total  :  {sum(engins.values())} XAF")
                total2.grid(row=1, column=0)
                ligne = 0
                for k, v in engins.items():
                    vue = customtkinter.CTkLabel(master=minor_on, text=f"{k}: ",font=customtkinter.CTkFont(size=14) )
                    vue.grid(row=ligne, column=0, padx=(10, 5), pady=(10,10))
                    label = customtkinter.CTkLabel(master=minor_on, text=f"{v}", font=customtkinter.CTkFont(size=14))
                    label.grid(row=ligne, column=2, padx=(10, 5))
                    quant = customtkinter.CTkLabel(master=minor_on, text=f"{quant2[ligne]}", font=customtkinter.CTkFont(size=14))
                    quant.grid(row=ligne, column=3,padx=(50, 5), sticky="e")
                    ligne += 1
        
        #frame pour les l'onglet materiaux
                sr = customtkinter.CTkFrame(master=textbox,  border_color="red", )
                sr.grid(row=0, column=2, sticky="nsew",pady=(5, 5), )
        
        #label
                sr_on = customtkinter.CTkFrame(master=sr,  border_width=1, height=50)
                sr_on.grid(row=0, column=0)
                cell_name_lb = customtkinter.CTkLabel(master=sr_on, text="Materiaux",width=350,  height=50,font=customtkinter.CTkFont(size=14, weight="bold"))
                cell_name_lb.grid(row=0, column=0, sticky="nsew")
    
                warning_on = customtkinter.CTkFrame(master=sr_on,  border_width=1, border_color="#DC3F3F")
                warning_on.grid(row=1, column=0, sticky="nsew",padx=(5, 5), pady=(5, 5))   
                
                total3 = customtkinter.CTkLabel(master=sr, text=f"Total  :  {sum(materiau.values())} XAF")
                total3.grid(row=1, column=0)
                ligne = 0
                for k, v in materiau.items():
                    vue = customtkinter.CTkLabel(master=warning_on, text=f"{k}: ",font=customtkinter.CTkFont(size=14) )
                    vue.grid(row=ligne, column=0, padx=(10, 5), pady=(10,10))
                    label = customtkinter.CTkLabel(master=warning_on, text=f"{v}", font=customtkinter.CTkFont(size=14))
                    label.grid(row=ligne, column=2, padx=(10, 5))
                    quant = customtkinter.CTkLabel(master=warning_on, text=f"{quant3[ligne]}", font=customtkinter.CTkFont(size=14))
                    quant.grid(row=ligne, column=3,padx=(50, 5), sticky="e")
                    ligne += 1
                    
                                
                total_main = sum(main_oeuvre.values())
                total_en = sum(engins.values())
                total_mat = sum(materiau.values())                
                bt_type = customtkinter.CTkLabel(master=textbox, text=f"Type de beton : {valeur}        Pour un Total de {sum([total_en, total_main, total_mat])}XAF")
                bt_type.grid(row=1, column=0)
                
                materiaux = pd.DataFrame(mat)
                materiaux.to_excel("materiaux.xlsx",index=False)
                
                engin = pd.DataFrame(eng)
                engin.to_excel("engins.xlsx",index=False)
                
                mains = pd.DataFrame(main)
                mains.to_excel("mains.xlsx",index=False)
                with open("sum.json", "w") as file:
                    json.dump(sum([total_en, total_main, total_mat]), file)
                with open("sum_mat.json", "w") as file:
                    json.dump(total_mat, file)
                with open("sum_eng.json", "w") as file:
                    json.dump(total_en, file)
                with open("sum_main.json", "w") as file:
                    json.dump(total_main, file)

                    
                    

            except Exception as e:
                new = customtkinter.CTkToplevel(self)
                new.title("Erreur de fichier")
                new.geometry("340x200")
                new.resizable(False, False)
                new.attributes('-topmost', True)
                new.grab_set()
                # new.grid_columnconfigure(0, weight=1)
                msg = customtkinter.CTkLabel(master=new, text="Une erreur c'est produit")
                msg.grid(row=0, column=0, padx=(100, 0), pady=(20, 40), sticky="news")
                bt1 = customtkinter.CTkButton(master=new, text="Fermer", command=lambda:local(), fg_color="red", hover_color="#DC3F3F")
                bt1.grid(row=1, column=0, padx=(100, 0), sticky="news")
                def local():
                    new.destroy()
                    
            return eng, mat, main
        
    def import_file(self):
        root = Tk()
        root.withdraw()
        path = filedialog.askopenfile(filetypes=[("Excel files",  ".xlsx .xls .csv")])
        self.file = path
        


    def open_input_dialog_event(self):
        dialog = customtkinter.CTkInputDialog(text="Type in a number:", title="CTkInputDialog")
        print("CTkInputDialog:", dialog.get_input())

    def change_appearance_mode_event(self, new_appearance_mode: str):
        customtkinter.set_appearance_mode(new_appearance_mode)

    def bti(self, choice,):
        global valeur 
        valeur = choice
        if self.file == "":
                new = customtkinter.CTkToplevel(self)
                new.title("Erreur de fichier")
                new.geometry("340x200")
                new.resizable(False, False)
                new.attributes('-topmost', True)
                new.grab_set()
                # new.grid_columnconfigure(0, weight=1)
                msg = customtkinter.CTkLabel(master=new, text="Aucun fichier n'a ete charger")
                msg.grid(row=0, column=0, padx=(100, 0), pady=(20, 40), sticky="news")
                bt1 = customtkinter.CTkButton(master=new, text="Fermer", command=lambda:local(), fg_color="red", hover_color="#DC3F3F")
                bt1.grid(row=1, column=0, padx=(100, 0), sticky="news")
                def local():
                    new.destroy()

        else:
            try:
                url = self.file.name
                bt = pd.read_excel(url)
                bti = bt[bt["Type de beton"] == f"{choice}"]
                self.cim.config(text=f"{bti.loc[0,"Dosage en ciment (KG/m3)"]}")
                self.gra.config(text=f"{bti.loc[0,"Dosage en gravier 00/15 (KG/m3)"]}")
                self.gra2.config(text=f"{bti.loc[0,"Dosage en gravier 15/25 (KG/m3)"]}")
                self.eau.config(text=f"{bti.loc[0,"Dosage en eau (L)"]}")
                self.sable.config(text=f"{bti.loc[0,"Dosage en sable (KG/m3)"]}")
                self.type.config(text=f"{choice}")
            except Exception as e:
                print(e)
        return choice

    def sidebar_button_event(self):
        print("sidebar_button click")
    def saves(self):
        # print(self.eng)
        pass
           
    def etat(self, choice):
        if os.path.exists("mains.xlsx") or os.path.exists("materiaux.xlsx") and os.path.exists("engins.xlsx"):
            try:
                print("chargement")
                mains = pd.read_excel("mains.xlsx")
                materiaux = pd.read_excel("materiaux.xlsx")
                engins = pd.read_excel("engins.xlsx")
                if choice == "PDF":
                    pdf = canvas.Canvas("BT.pdf",pagesize=letter)
                    largeur, hauteur = letter
                    pdf.drawString(100, hauteur - 100, "TYPE DE BETON: ")
                    pdf.rect(50, hauteur - 150, 200, 100)
                    # root = Tk()
                    # root.withdraw()
                    # path = filedialog.askdirectory()
                    pdf.save()
                    pass
                
                elif choice == "Exel":
                    # data = pd.read_excel("etat.xlsx")
                    root = Tk()
                    root.withdraw()
                    path = filedialog.askdirectory()
                    # data.to_excel(f"{path}/BT.xlsx")
                elif choice == "Word":

                    doc = Document()
                    doc.add_heading(f"Type de BT   {valeur}  Pour un Total de")
                    
                    doc.add_heading("Main d'oeuvre")
                    tab = doc.add_table(rows=1, cols=len(mains.columns), style="Table Grid")
                    

                    
                    hdr_cells = tab.rows[0].cells 
                    for i, column_name in enumerate(mains.columns): 
                        hdr_cells[i].text = column_name
                        
                    for index, row in mains.iterrows(): 
                        row_cells = tab.add_row().cells 
                        for i, column_name in enumerate(mains.columns): 
                            row_cells[i].text = str(row[column_name])
                        ################## 
                    doc.add_page_break()
                    doc.add_paragraph("")    
                    doc.add_heading("Engins")  
                     
                    tab2 = doc.add_table(rows=2, cols=len(materiaux.columns), style="Table Grid")
                    
                    
                    
                    
                    hdr_cells2 = tab2.rows[0].cells 
                    for i, column_name1 in enumerate(materiaux.columns): 
                        hdr_cells2[i].text = column_name1
                        
                    for index, row in materiaux.iterrows(): 
                        row_cells = tab2.add_row().cells 
                        for i, column_name1 in enumerate(materiaux.columns): 
                            row_cells[i].text = str(row[column_name1])
                            
                    doc.add_page_break()
                    doc.add_paragraph("")        
                    doc.add_heading("Materiaux")  
                     
                    tab3 = doc.add_table(rows=3, cols=len(engins.columns), style="Table Grid")
                    
                    
                    
                    
                    hdr_cells3 = tab3.rows[0].cells 
                    for i, column_name2 in enumerate(engins.columns): 
                        hdr_cells3[i].text = column_name2
                        
                    for index, row in engins.iterrows(): 
                        row_cells = tab3.add_row().cells 
                        for i, column_name2 in enumerate(engins.columns): 
                            row_cells[i].text = str(row[column_name2])
                    
                    root = Tk()
                    root.withdraw()
                    path = filedialog.askdirectory()
                    doc.save(f"{path}/BT.docx")
                
            except Exception as e:
                print(e)
                
        else:
            print("chemin non touver")


if __name__ == "__main__":
    app = App()
    app.mainloop()

